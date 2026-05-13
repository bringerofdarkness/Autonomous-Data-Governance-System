import csv
import json

from pathlib import Path
from typing import Any

from datetime import date, datetime

from openpyxl import load_workbook

from docx import Document
from pypdf import PdfReader

from app.core.config import get_settings


settings = get_settings()

SUPPORTED_TEXT_EXTENSIONS = {".txt"}
SUPPORTED_PDF_EXTENSIONS = {".pdf"}
SUPPORTED_DOCX_EXTENSIONS = {".docx"}
SUPPORTED_CSV_EXTENSIONS = {".csv"}
SUPPORTED_XLSX_EXTENSIONS = {".xlsx"}
SUPPORTED_JSON_EXTENSIONS = {".json"}


def _extract_txt_text(file_path: Path) -> dict[str, Any]:
    try:
        extracted_text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError("Only UTF-8 text files are currently supported.") from exc

    return {
        "extraction_method": "utf8_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
        },
        "warnings": [],
    }


def _extract_pdf_text(file_path: Path) -> dict[str, Any]:
    reader = PdfReader(str(file_path))

    page_texts: list[str] = []
    warnings: list[str] = []

    for page_index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""

        if not page_text.strip():
            warnings.append(f"No extractable text found on page {page_index + 1}.")

        page_texts.append(page_text)

    extracted_text = "\n\n".join(page_texts).strip()

    if not extracted_text:
        raise ValueError(
            "No extractable text found in PDF. Scanned PDFs will require OCR support later."
        )

    return {
        "extraction_method": "pypdf_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
            "page_count": len(reader.pages),
        },
        "warnings": warnings,
    }


def _extract_docx_text(file_path: Path) -> dict[str, Any]:
    document = Document(str(file_path))

    paragraph_texts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    table_texts: list[str] = []

    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]

            if row_values:
                table_texts.append(" | ".join(row_values))

    extracted_text = "\n".join(paragraph_texts + table_texts).strip()

    if not extracted_text:
        raise ValueError("No extractable text found in DOCX document.")

    return {
        "extraction_method": "python_docx_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
            "paragraph_count": len(paragraph_texts),
            "table_count": len(document.tables),
        },
        "warnings": [],
    }

def _extract_csv_text(file_path: Path) -> dict[str, Any]:
    rows: list[list[str]] = []

    try:
       with file_path.open("r", encoding="utf-8-sig", newline="") as csv_file:
            reader = csv.reader(csv_file)

            for row in reader:
                cleaned_row = [
                    cell.strip()
                    for cell in row
                ]
                rows.append(cleaned_row)

    except UnicodeDecodeError as exc:
        raise ValueError("Only UTF-8 or UTF-8-BOM CSV files are currently supported.") from exc

    if not rows:
        raise ValueError("No rows found in CSV file.")

    header = rows[0]
    data_rows = rows[1:]

    text_lines = [
        "ADGS CSV Extracted Document",
        f"Columns: {', '.join(header)}",
        "",
        "Rows:",
    ]

    for row_index, row in enumerate(data_rows, start=1):
        row_pairs = []

        for column_index, cell in enumerate(row):
            column_name = (
                header[column_index]
                if column_index < len(header)
                else f"column_{column_index + 1}"
            )
            row_pairs.append(f"{column_name}: {cell}")

        text_lines.append(f"Row {row_index}: " + " | ".join(row_pairs))

    extracted_text = "\n".join(text_lines).strip()

    return {
        "extraction_method": "csv_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
            "row_count": len(data_rows),
            "column_count": len(header),
            "columns": header,
        },
        "warnings": [],
    }


def _excel_cell_to_text(value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, (datetime, date)):
        return value.isoformat()

    return str(value).strip()


def _extract_xlsx_text(file_path: Path) -> dict[str, Any]:
    workbook = load_workbook(
        filename=str(file_path),
        read_only=True,
        data_only=True,
    )

    text_lines: list[str] = ["ADGS XLSX Extracted Document"]
    warnings: list[str] = []
    sheet_metadata: list[dict[str, Any]] = []

    total_data_rows = 0

    for worksheet in workbook.worksheets:
        sheet_name = worksheet.title

        rows = [
            [
                _excel_cell_to_text(cell)
                for cell in row
            ]
            for row in worksheet.iter_rows(values_only=True)
        ]

        non_empty_rows = [
            row
            for row in rows
            if any(cell for cell in row)
        ]

        if not non_empty_rows:
            warnings.append(f"No extractable data found in sheet '{sheet_name}'.")
            sheet_metadata.append(
                {
                    "sheet_name": sheet_name,
                    "row_count": 0,
                    "column_count": 0,
                    "columns": [],
                }
            )
            continue

        header = non_empty_rows[0]
        data_rows = non_empty_rows[1:]

        normalized_header = [
            column_name if column_name else f"column_{index + 1}"
            for index, column_name in enumerate(header)
        ]

        text_lines.append("")
        text_lines.append(f"Sheet: {sheet_name}")
        text_lines.append(f"Columns: {', '.join(normalized_header)}")
        text_lines.append("Rows:")

        for row_index, row in enumerate(data_rows, start=1):
            row_pairs = []

            for column_index, cell in enumerate(row):
                column_name = (
                    normalized_header[column_index]
                    if column_index < len(normalized_header)
                    else f"column_{column_index + 1}"
                )

                row_pairs.append(f"{column_name}: {cell}")

            text_lines.append(f"Row {row_index}: " + " | ".join(row_pairs))

        total_data_rows += len(data_rows)

        sheet_metadata.append(
            {
                "sheet_name": sheet_name,
                "row_count": len(data_rows),
                "column_count": len(normalized_header),
                "columns": normalized_header,
            }
        )

    workbook.close()

    extracted_text = "\n".join(text_lines).strip()

    if not extracted_text or extracted_text == "ADGS XLSX Extracted Document":
        raise ValueError("No extractable text found in XLSX file.")

    return {
        "extraction_method": "openpyxl_xlsx_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
            "sheet_count": len(sheet_metadata),
            "total_data_rows": total_data_rows,
            "sheets": sheet_metadata,
        },
        "warnings": warnings,
    }

def _flatten_json_value(value: Any, prefix: str = "") -> list[str]:
    lines: list[str] = []

    if isinstance(value, dict):
        for key, item in value.items():
            next_prefix = f"{prefix}.{key}" if prefix else str(key)
            lines.extend(_flatten_json_value(item, next_prefix))

    elif isinstance(value, list):
        for index, item in enumerate(value):
            next_prefix = f"{prefix}[{index}]"
            lines.extend(_flatten_json_value(item, next_prefix))

    else:
        display_value = "" if value is None else str(value)
        lines.append(f"{prefix}: {display_value}")

    return lines


def _count_json_keys(value: Any) -> int:
    if isinstance(value, dict):
        return len(value) + sum(_count_json_keys(item) for item in value.values())

    if isinstance(value, list):
        return sum(_count_json_keys(item) for item in value)

    return 0


def _get_json_max_depth(value: Any, current_depth: int = 0) -> int:
    if isinstance(value, dict):
        if not value:
            return current_depth

        return max(
            _get_json_max_depth(item, current_depth + 1)
            for item in value.values()
        )

    if isinstance(value, list):
        if not value:
            return current_depth

        return max(
            _get_json_max_depth(item, current_depth + 1)
            for item in value
        )

    return current_depth


def _extract_json_text(file_path: Path) -> dict[str, Any]:
    try:
        with file_path.open("r", encoding="utf-8-sig") as json_file:
            json_data = json.load(json_file)

    except UnicodeDecodeError as exc:
        raise ValueError("Only UTF-8 or UTF-8-BOM JSON files are currently supported.") from exc

    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON file: {exc}") from exc

    flattened_lines = _flatten_json_value(json_data)

    if not flattened_lines:
        raise ValueError("No extractable data found in JSON file.")

    text_lines = [
        "ADGS JSON Extracted Document",
        "",
        "Fields:",
        *flattened_lines,
    ]

    extracted_text = "\n".join(text_lines).strip()

    return {
        "extraction_method": "json_text",
        "text": extracted_text,
        "metadata": {
            "char_count": len(extracted_text),
            "file_size_bytes": file_path.stat().st_size,
            "top_level_type": type(json_data).__name__,
            "key_count": _count_json_keys(json_data),
            "max_depth": _get_json_max_depth(json_data),
        },
        "warnings": [],
    }


def extract_document_text(stored_filename: str) -> dict[str, Any]:
    if not stored_filename:
        raise ValueError("stored_filename is missing.")

    file_path = Path(settings.UPLOAD_DIR) / stored_filename

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_extension = file_path.suffix.lower()

    if file_extension in SUPPORTED_TEXT_EXTENSIONS:
        extraction_result = _extract_txt_text(file_path)
    elif file_extension in SUPPORTED_PDF_EXTENSIONS:
        extraction_result = _extract_pdf_text(file_path)
    elif file_extension in SUPPORTED_DOCX_EXTENSIONS:
        extraction_result = _extract_docx_text(file_path)
    elif file_extension in SUPPORTED_CSV_EXTENSIONS:
        extraction_result = _extract_csv_text(file_path)
    elif file_extension in SUPPORTED_XLSX_EXTENSIONS:
        extraction_result = _extract_xlsx_text(file_path)
    elif file_extension in SUPPORTED_JSON_EXTENSIONS:
        extraction_result = _extract_json_text(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{file_extension}'. Currently supported: .txt, .pdf, .docx, .csv, .xlsx, .json"
        )

    return {
        "stored_filename": stored_filename,
        "file_path": str(file_path),
        "file_extension": file_extension,
        "extraction_method": extraction_result["extraction_method"],
        "text": extraction_result["text"],
        "metadata": extraction_result["metadata"],
        "warnings": extraction_result["warnings"],
    }